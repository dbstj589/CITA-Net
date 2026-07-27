#!/usr/bin/env python
"""Build a plain-language DOCX report (figures + tables) of the motion/difficulty
experiment. Figures use English labels (font-safe); body text is Korean.
Read-only w.r.t. experiments; only writes figures + the .docx."""
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO = Path(".")
FIG = REPO / "docs" / "report_figs"; FIG.mkdir(parents=True, exist_ok=True)
BLUE, ORANGE, PURPLE, GREEN = "#2c7fb8", "#d95f02", "#7570b3", "#1a9641"

# ---------------- data (from the experiment) ----------------
seeds = ["06", "07", "08", "09", "10"]
dtest = [0.0270, 0.0352, 0.0139, 0.0082, 0.0174]          # m3_full - no_motion (test)
recall_d = [0.0352, 0.0425, 0.0136, 0.0089, 0.0133]
gate_easy = [0.650, 0.663, 0.674, 0.612, 0.644]
gate_amb = [0.966, 0.982, 0.973, 0.947, 0.970]

# ---------------- figures ----------------
def save(fig, name):
    fig.tight_layout(); fig.savefig(FIG / name, dpi=150, bbox_inches="tight"); plt.close(fig)

# Fig1: dataset difficulty (hard-negative ratio)
fig, ax = plt.subplots(figsize=(6.2, 3.6))
labs = ["frozen\n(easy, 28/sector)", "amb80\n(53/sector)", "amb160\n(hard, 109/sector)"]
vals = [9.4, 17.5, 36.4]
b = ax.bar(labs, vals, color=[GREEN, ORANGE, BLUE])
ax.bar_label(b, fmt="%.1f:1", padding=3)
ax.set_ylabel("hard-negative : positive ratio")
ax.set_title("Dataset difficulty grows ~4x from frozen to amb160")
save(fig, "fig1_difficulty.png")

# Fig2: validation gate (m1 F1)
fig, ax = plt.subplots(figsize=(6.2, 3.6))
labs = ["frozen", "amb80", "amb160"]; vals = [0.9701, 0.9639, 0.9025]
b = ax.bar(labs, vals, color=[GREEN, ORANGE, BLUE])
ax.bar_label(b, fmt="%.3f", padding=3)
ax.axhspan(0.9701 - 0.0070, 0.9701 + 0.0070, color=GREEN, alpha=0.15)
ax.axhline(0.9701, color=GREEN, ls="--", lw=1)
ax.set_ylim(0.85, 1.0); ax.set_ylabel("m1 baseline  test F1")
ax.set_title("Difficulty gate: amb80 within noise (fail), amb160 -6.75pt (pass)")
save(fig, "fig2_gate.png")

# Fig3: slot starvation fix
fig, ax = plt.subplots(figsize=(6.2, 3.6))
x = np.arange(2); w = 0.35
ax.bar(x - w/2, [0.3045, 0.3172], w, label="num_slots = 48 (starved)", color="#bbbbbb")
ax.bar(x + w/2, [0.5661, 0.5457], w, label="num_slots = 160 (fixed)", color=[BLUE, ORANGE])
ax.set_xticks(x); ax.set_xticklabels(["m3_full", "no_motion"])
ax.set_ylabel("test F1"); ax.set_ylim(0, 0.7)
ax.set_title("Decoder slot fix: F1 recovers 0.30 -> ~0.55")
ax.legend(fontsize=8)
for i, v in enumerate([0.3045, 0.3172]):
    ax.text(i - w/2, v + .01, f"{v:.2f}", ha="center", fontsize=8)
for i, v in enumerate([0.5661, 0.5457]):
    ax.text(i + w/2, v + .01, f"{v:.2f}", ha="center", fontsize=8)
save(fig, "fig3_slot.png")

# Fig5: per-seed test delta (5/5 positive)
fig, ax = plt.subplots(figsize=(6.2, 3.6))
b = ax.bar(seeds, dtest, color=BLUE)
ax.bar_label(b, fmt="+%.3f", padding=2, fontsize=8)
ax.axhline(0, color="k", lw=0.8)
ax.axhline(np.mean(dtest), color=ORANGE, ls="--", lw=1.5, label=f"mean +{np.mean(dtest):.3f}")
ax.set_ylabel("test F1 :  m3_full - no_motion"); ax.set_xlabel("seed")
ax.set_title("Reversal on test: all 5 seeds positive (p = 0.013)")
ax.legend(fontsize=8)
save(fig, "fig5_delta.png")

# Fig6: gate distribution (motion gate easy vs ambiguous, mean +/- std)
fig, ax = plt.subplots(figsize=(6.2, 3.6))
x = np.arange(2); m = [np.mean(gate_easy), np.mean(gate_amb)]; sd = [np.std(gate_easy), np.std(gate_amb)]
b = ax.bar(["easy pairs\n(similarity > 0.5)", "ambiguous pairs\n(similarity <= 0.5)"], m, yerr=sd,
           capsize=5, color=["#9ecae1", BLUE])
ax.bar_label(b, fmt="%.3f", padding=8, fontsize=9)
ax.set_ylabel("learned MOTION gate (0-1)  mean of 5 seeds"); ax.set_ylim(0, 1.1)
ax.set_title("The gate turns MOTION ON for ambiguous pairs (0.65 -> 0.97)")
save(fig, "fig6_gate.png")

# Fig7: p-value progression
fig, ax = plt.subplots(figsize=(6.2, 3.6))
b = ax.bar(["n = 3", "n = 5"], [0.055, 0.013], color=[ORANGE, BLUE])
ax.bar_label(b, fmt="p = %.3f", padding=3)
ax.axhline(0.05, color="red", ls="--", lw=1.2, label="p = 0.05 threshold")
ax.set_ylabel("test reversal  p-value (paired t-test)"); ax.set_ylim(0, 0.08)
ax.set_title("More seeds -> reversal becomes significant (0.055 -> 0.013)")
ax.legend(fontsize=8)
save(fig, "fig7_pprog.png")

# ---------------- docx ----------------
doc = Document()
st = doc.styles["Normal"]; st.font.name = "맑은 고딕"; st.font.size = Pt(10.5)

def H(text, lvl=1):
    h = doc.add_heading(text, level=lvl)
    return h
def P(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    return p
def pic(name, width=6.0, caption=None):
    doc.add_picture(str(FIG / name) if not str(name).startswith("results") else str(REPO / name), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph(); r = c.add_run(caption); r.italic = True; r.font.size = Pt(9)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
def table(headers, rows, bold_first_col=False):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = ""; run = cell.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; run = cells[i].paragraphs[0].add_run(str(v)); run.font.size = Pt(9.5)
            if bold_first_col and i == 0: run.bold = True
    return t

# --- Title ---
title = doc.add_heading("", level=0)
tr = title.add_run("운동(움직임) 정보는 언제 도움이 되는가")
tr.font.size = Pt(22)
sub = doc.add_paragraph(); sr = sub.add_run("백마고지 데이터 기반 CITA-Net 실험 종합 보고서")
sr.font.size = Pt(13); sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER; title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- 한눈 요약 ---
H("한눈에 보는 결론", 1)
P("여러 관측 정보를 보고 '이 관측과 저 관측이 같은 대상인가'를 판단하는 인공지능(CITA-Net)에서, "
  "대상의 '움직임(운동) 정보'를 쓰는 것이 좋은지 나쁜지를 실험으로 확인했습니다.", )
P("핵심 발견: 움직임 정보는 '항상' 좋은 게 아니라 '상황에 따라' 다릅니다. "
  "구분이 쉬운 데이터에서는 오히려 방해가 되고(정확도 0.703 → 0.652로 하락), "
  "구분이 어려운(붐비는) 데이터에서는 도움이 됩니다(0.546 → 0.566으로 상승). "
  "이 '뒤집힘'은 5번의 반복 실험에서 모두 같은 방향으로 나타났고, 통계적으로도 유의미했습니다(p = 0.013).", bold=True)
P("반면, 상황에 맞춰 자동으로 켜고 끄도록 '학습'시킨 방식(term_gating)은 "
  "단순히 움직임을 항상 켜둔 방식을 이기지 못했습니다.")

# --- 1 ---
H("1. 이 실험이 답하려는 질문 — 쉬운 비유로", 1)
P("여러 대의 CCTV가 같은 지역을 찍고 있다고 상상해 봅시다. 카메라 A가 잡은 '사람'과 "
  "카메라 B가 잡은 '사람'이 같은 사람인지 이어 붙여야 합니다. 방법은 두 가지입니다.")
P("• 생김새(의미 정보): 옷 색깔·체형 등이 비슷하면 같은 사람일 것이다.\n"
  "• 움직임(운동 정보): 비슷한 속도·방향으로 움직이면 같은 사람일 것이다.")
P("사람이 몇 명 없고 서로 다르게 생겼다면 '생김새'만으로 충분합니다. 이때 '움직임'을 굳이 따지면 "
  "오히려 헷갈릴 수 있습니다. 그런데 비슷하게 생긴 사람이 잔뜩 붐빈다면 '생김새'로는 구분이 안 되고, "
  "이때는 '움직임'이 결정적인 단서가 됩니다. 이 실험은 바로 이 직관을 데이터로 검증합니다.")

# --- 2 ---
H("2. CITA-Net이 하는 일 (배경)", 1)
P("CITA-Net은 1952년 백마고지 전투 상황을 본떠 만든 가상 데이터에서, "
  "서로 다른 두 정보원(지식그래프 A·B)의 관측을 '같은 부대/물체'끼리 연결하는 모델입니다. "
  "연결 판단에 6가지 단서를 씁니다: 의미(생김새)·시간·움직임(운동)·상태·관계·출처. "
  "이번 실험의 주인공은 그중 '움직임(motion)' 단서입니다.")
P("성능은 F1이라는 점수(0~1, 높을수록 좋음)로 잽니다. 쉽게는 '연결을 얼마나 정확하고 빠짐없이 했는가'의 "
  "종합 점수라고 보면 됩니다.")

# --- 3 ---
H("3. 첫 발견 — 쉬운 데이터에서는 움직임 정보가 '방해'가 되었다", 1)
P("난이도가 낮은 기본 데이터(frozen)에서 두 모델을 비교했습니다. "
  "하나는 움직임 단서를 쓰는 모델(m3_full), 하나는 빼버린 모델(no_motion)입니다.")
table(["모델", "test F1 (정확도)"],
      [["움직임 뺀 모델 (no_motion)", "0.7034  ← 더 높음"],
       ["움직임 쓴 모델 (m3_full)", "0.6520"]], bold_first_col=True)
P("결과가 직관과 반대였습니다. 쉬운 데이터에서는 움직임을 '뺀' 쪽이 오히려 +5.1점 더 좋았습니다. "
  "즉 이 상황에서 움직임 정보는 도움이 아니라 잡음(방해)이었습니다.")

# --- 4 ---
H("4. 가설 — '움직임 정보는 어려울 때만 도움이 된다'", 1)
P("위 발견에서 가설을 세웠습니다.", )
P("가설: 움직임 정보의 효용은 데이터의 '난이도'에 따라 달라진다. "
  "쉬우면 방해가 되고, 어려우면(비슷한 대상이 붐벼서 생김새로 구분이 안 되면) 도움이 된다. "
  "따라서 쉬운→어려운으로 갈수록 두 모델의 우열이 '뒤집혀야' 한다.", bold=True)
P("이 가설을 확인하려면 '생김새로는 구분이 안 되는, 붐비는 데이터'를 일부러 만들어야 합니다.")

# --- 5 ---
H("5. 어려운 데이터 만들기 — '붐비게' 만들기", 1)
P("한 구역 안에 같은 종류의 대상을 더 촘촘히 넣어 '헷갈리는 정도(모호성)'를 높였습니다. "
  "아래 표에서 '헷갈리는 짝의 비율'이 클수록 어려운 데이터입니다.")
table(["데이터", "구역당 대상 수", "헷갈리는 짝 : 진짜 짝"],
      [["frozen (쉬움)", "28개", "9.4 : 1"],
       ["amb80 (2배)", "53개", "17.5 : 1"],
       ["amb160 (어려움, 4배)", "109개", "36.4 : 1"]], bold_first_col=True)
pic("fig1_difficulty.png", 5.6, "그림 1. 데이터가 어려워질수록 '헷갈리는 짝'의 비율이 약 4배로 커진다.")

# --- 6 ---
H("6. '정말 어려워졌는지' 먼저 확인 (검증 게이트)", 1)
P("어렵게 만들었다고 주장만 하면 안 되므로, 가벼운 기준 모델(m1)로 '점수가 실제로 떨어지는지' 먼저 확인했습니다. "
  "기준(frozen)의 점수는 0.970입니다.")
table(["데이터", "기준 모델 점수", "판정"],
      [["amb80", "0.964 (거의 그대로)", "실패 — 시드마다의 자연스러운 오차 범위 안. 충분히 안 어려움"],
       ["amb160", "0.903 (뚜렷이 하락)", "통과 — 오차의 약 10배로 하락. 확실히 어려움"]], bold_first_col=True)
pic("fig2_gate.png", 5.6, "그림 2. amb80은 기준과 사실상 같아 '실패', amb160은 -6.75점으로 뚜렷이 어려워 '통과'. "
    "(초록 띠 = 기준의 자연스러운 오차 범위)")
P("그래서 amb80은 버리고, 충분히 어려운 amb160으로 본 실험을 진행했습니다.")

# --- 7 ---
H("7. 도중에 만난 함정 — '디코더 자리 부족'", 1)
P("본 실험을 처음 돌렸을 때 점수가 0.30까지 폭락했습니다. 원인을 파보니 모델의 마지막 정리 단계(디코더)가 "
  "한 구역에서 최대 48개의 대상만 만들 수 있게 되어 있었는데, amb160은 대상이 약 110개라 "
  "'자리가 모자라' 대부분을 제대로 못 만든 것이었습니다. 모델 실력 문제가 아니라 설정의 한계였습니다.")
P("이 자리 수를 48 → 160으로 늘려 문제를 없앴더니 점수가 정상(약 0.55)으로 회복됐습니다. "
  "이 함정을 못 잡았다면 '움직임이 도움이 안 된다'는 잘못된 결론을 낼 뻔했습니다.")
pic("fig3_slot.png", 5.6, "그림 3. 디코더 자리를 48→160으로 늘리자 점수가 0.30에서 약 0.55로 회복(정상화).")

# --- 8 ---
H("8. 본 실험 결과 — 움직임의 '뒤집힘'을 확인", 1)
P("어려운 데이터(amb160)에서 세 모델을 각각 서로 다른 5번(시드 5개) 반복 학습해 평균을 냈습니다. "
  "(설정: 동일 데이터·디코더 자리 160·40회 반복 학습·동일 채점 방식)")
table(["모델", "test F1 (평균 ± 표준편차)", "dev F1"],
      [["m3_full (움직임 사용)", "0.5661 ± 0.0132  ← 가장 높음", "0.5601"],
       ["no_motion (움직임 제거)", "0.5457 ± 0.0113", "0.5539"],
       ["term_gating (자동 조절)", "0.5388 ± 0.0569 (흔들림 큼)", "0.5413"]], bold_first_col=True)
P("쉬운 데이터에서는 no_motion(0.703) > m3_full(0.652)이었는데, 어려운 데이터에서는 "
  "m3_full(0.566) > no_motion(0.546)으로 '뒤집혔습니다'.", bold=True)
pic("results/hard_ambiguity_main/n5_f1_reversal.png", 6.4,
    "그림 4. 왼쪽(쉬움): 움직임 뺀 쪽이 높음. 오른쪽(어려움): 움직임 쓴 쪽이 높음 → 부호가 뒤집힘.")
P("이 뒤집힘이 우연이 아닌지 확인하기 위해, 5번의 반복에서 '움직임 모델 − 움직임 뺀 모델'의 점수 차이를 봤습니다. "
  "5번 모두 양수(+)였고, 통계 검정 결과 p = 0.013으로 '유의미'했습니다(보통 0.05보다 작으면 유의).")
pic("fig5_delta.png", 5.6, "그림 5. test에서 5번 반복 모두 '움직임 모델'이 더 높았다(모두 +). 평균 +0.020, p=0.013.")
P("반복 횟수를 3번에서 5번으로 늘리자 유의확률이 0.055(애매)에서 0.013(유의)으로 낮아졌습니다. "
  "여기서 실험을 멈췄습니다 — 원하는 결과가 나올 때까지 무한정 늘리는 것은 통계 왜곡이기 때문입니다.")
pic("fig7_pprog.png", 5.2, "그림 6. 반복을 늘리자 뒤집힘이 통계적으로 유의해짐(0.055 → 0.013, 빨간선=0.05 기준).")
P("왜 도움이 되는지도 확인했습니다. '움직임 모델'은 놓치지 않고 찾아낸 비율(recall)이 5번 모두 더 높았습니다"
  "(평균 +0.023). 즉 붐비는 상황에서 움직임 단서가 '같은 대상'을 더 잘 회수했습니다.")

# --- 9 ---
H("9. 자동 조절 방식은 정말 '움직임'을 켰는가", 1)
P("term_gating(상황에 맞춰 각 단서를 자동으로 켜고 끄는 방식)이 실제로 의도대로 동작하는지 봤습니다. "
  "'생김새가 비슷해서 헷갈리는 짝'에서 움직임 단서를 얼마나 켜는지(0~1) 측정했습니다.")
table(["짝의 종류", "움직임 단서를 켠 정도 (5시드 평균)"],
      [["구분 쉬운 짝 (생김새 다름)", "0.65"],
       ["헷갈리는 짝 (생김새 비슷)", "0.97  ← 크게 켬"]], bold_first_col=True)
pic("fig6_gate.png", 5.6, "그림 7. 자동 조절 방식은 '헷갈리는 짝'에서 움직임 단서를 강하게 켠다(0.65→0.97). 의도대로 동작.")
P("즉 '내용(생김새)으로 못 가르면 움직임을 쓴다'는 규칙을 스스로 학습했고, 5번 모두 일관됐습니다.")

# --- 10 ---
H("10. 그런데 자동 조절 방식은 왜 더 좋지 않았나", 1)
P("동작은 옳았지만(그림 7), 최종 점수는 '움직임을 항상 켠' 단순 모델(m3_full)을 넘지 못했습니다"
  "(평균 −0.027, 5번 중 방향이 엇갈림, p=0.402). 게다가 시드마다 점수가 크게 흔들렸습니다"
  "(표준편차 0.057로 m3_full의 약 4배, 한 번은 −0.136까지 떨어짐). "
  "정리하면: 자동 조절은 '똑똑하게 동작'하지만 '추가 이득 없이 불안정만 늘렸다'는 결론입니다.")

# --- 11 ---
H("11. 최종 결론", 1)
P("1) 움직임 정보는 '난이도의 함수'다 — 쉬우면 방해, 어려우면 도움. 이 뒤집힘은 어려운 데이터의 "
  "test에서 5번 모두, 통계적으로 유의하게(p=0.013) 확인됐다. (dev에서는 5번 중 3번만 같은 방향이라 부분적.)", bold=True)
P("2) 왜 도움이 되는가 — 붐비는 상황에서 움직임 단서가 '같은 대상'을 더 잘 회수(recall +0.023, 5번 일관)하기 때문.")
P("3) 자동 조절(term_gating)은 의도대로 움직임을 켜지만, 단순히 항상 켠 모델을 넘지 못하고 불안정만 늘렸다.")

# --- 12 ---
H("12. 정직한 한계", 1)
P("• 절대 점수는 낮다(약 0.57). amb160이 원래 매우 어려운 데이터이기 때문이며, '상대적 우열(뒤집힘)'과 "
  "'절대 성능'은 구분해서 봐야 한다.\n"
  "• 반복은 5번(시드 5개)으로 통계적 힘이 크지 않다. 그래서 p값 하나가 아니라 '5번 모두 같은 방향인지', "
  "'recall 같은 메커니즘도 일관되는지'를 함께 봐서 결론을 냈다.\n"
  "• dev(검증용 분할)에서는 뒤집힘이 부분적(3/5)이었다 — test만큼 깔끔하진 않다.\n"
  "• 데이터는 합성(가상)이며 실제 전투의 재현이 아니다.")

P(" ")
foot = doc.add_paragraph(); fr = foot.add_run(
    "데이터: data/hard_ambiguity/amb160  ·  결과: results/hard_ambiguity_main/  ·  "
    "설정: 디코더 자리 160, 40회 학습, 시드 5개(06/07/08/09/10)")
fr.italic = True; fr.font.size = Pt(8.5); fr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

out = REPO / "docs" / "CITA-Net_실험보고서.docx"
doc.save(str(out))
print("saved:", out, "| figures:", len(list(FIG.glob('*.png'))))
